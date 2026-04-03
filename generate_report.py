#!/usr/bin/env python3
"""
Generate the TriML Final Report as a .docx file using python-docx.
CS 6140: Machine Learning — Northeastern University
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = os.path.dirname(os.path.abspath(__file__))
PLOTS = os.path.join(BASE, "results", "plots")
OUT = os.path.join(BASE, "TriML_Final_Report.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.name = "Times New Roman"
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = "Times New Roman"

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def add_figure(doc, path, caption, width=Inches(5.5)):
    """Insert an image with a caption paragraph below it."""
    if not os.path.exists(path):
        p = doc.add_paragraph(f"[Image not found: {path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else None
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run.italic = True


def body(doc, text):
    """Add a body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
    return p


def heading(doc, text, level=1):
    """Add a heading and style it."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
    return h


# ---------------------------------------------------------------------------
# Build Document
# ---------------------------------------------------------------------------
doc = Document()

# -- Default font --
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# -- Narrow margins --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ===== TITLE =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(4)
run = title_p.add_run("Predicting Athlete Training State and Injury Risk\nfrom Wearable Sensor Data")
run.bold = True
run.font.size = Pt(16)
run.font.name = "Times New Roman"

# ===== AUTHOR =====
author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_after = Pt(2)
run = author_p.add_run("William Felipe Quiroz")
run.font.size = Pt(11)
run.font.name = "Times New Roman"

affil_p = doc.add_paragraph()
affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil_p.paragraph_format.space_after = Pt(14)
run = affil_p.add_run("CS 6140: Machine Learning | Northeastern University")
run.font.size = Pt(11)
run.font.name = "Times New Roman"
run.italic = True

# ===== 1. ABSTRACT =====
heading(doc, "1. Abstract", level=1)

body(doc,
     "Overuse injuries in endurance athletes arise from accumulated physiological stress that "
     "exceeds the body's ability to recover, a process measurable through wearable sensor data. "
     "Rather than predicting injury as a future binary event (which is confounded by unknown "
     "training decisions), this project frames the problem as predicting an athlete's current "
     "training state from daily physiological signals. We engineer a composite Grit Score per day "
     "from training load and recovery data, then use it to (1) classify athletes into three "
     "readiness states\u2014Overreaching, Balanced, or Undertrained, (2) detect injury risk as a "
     "binary classification task, and (3) regress on the continuous Grit Score. Using the "
     "Synthetic Triathlete Dataset (Zenodo, 2025), which contains 366,000 daily records and "
     "384,153 activity sessions across 1,000 athletes, we compare three classifiers (Logistic "
     "Regression, Random Forest, DNN) and three regressors (Lasso with polynomial features, "
     "Random Forest, DNN), with systematic hyperparameter tuning via 5-fold GroupKFold "
     "cross-validation. The DNN achieves the best performance across all tasks: AUC-ROC 0.948 "
     "for injury classification, 0.997 for load-class classification, and R\u00b2 = 0.988 for "
     "Grit Score regression.")

# ===== 2. INTRODUCTION =====
heading(doc, "2. Introduction", level=1)

body(doc,
     "Endurance athletes\u2014triathletes, runners, cyclists, and swimmers\u2014face a persistent "
     "tension between training hard enough to improve and training so hard that the body breaks "
     "down. Overuse injuries account for the majority of training-related injuries in endurance "
     "sports, and they are fundamentally a problem of load management: when the cumulative "
     "physiological stress of recent training exceeds the body\u2019s capacity to recover, soft "
     "tissue damage accumulates until it crosses a clinical threshold. The acute-to-chronic "
     "workload ratio (ACWR), introduced by Gabbett (2016) and refined by Hulin et al. (2016), "
     "formalizes this concept by comparing short-term training load (typically 7 days) to a "
     "longer baseline (typically 28 days). An ACWR substantially above 1.0 indicates a rapid "
     "spike in load relative to what the athlete is accustomed to, and is associated with elevated "
     "injury risk.")

body(doc,
     "Modern wearable devices\u2014such as Garmin watches, WHOOP straps, and Oura rings\u2014now "
     "capture a rich set of physiological signals on a daily basis. Heart-rate variability (HRV), "
     "resting heart rate (RHR), sleep duration and quality metrics, stress scores, and body "
     "battery estimates provide a window into the autonomic nervous system\u2019s response to "
     "training and life stress. These signals reflect real-time recovery status: a suppressed HRV "
     "or elevated RHR relative to personal baselines indicates that the body is under strain, "
     "while robust sleep metrics suggest adequate recovery. Platforms like WHOOP and Garmin "
     "already surface daily readiness scores derived from these signals, but the underlying "
     "algorithms are proprietary and their predictive validity is not well-studied.")

body(doc,
     "Most existing approaches to injury prediction in the literature treat the problem as a "
     "binary future-event prediction task: given today\u2019s data, will the athlete be injured in "
     "the next N days? This framing is inherently limited because the outcome depends on future "
     "training decisions that are unknown at prediction time. If a model predicts high injury risk "
     "and the athlete reduces training in response, the injury may never occur\u2014making the "
     "prediction appear incorrect even though it was clinically useful. Simple threshold-based "
     "rules (e.g., flag an athlete whenever ACWR exceeds 1.5) are also limited because they "
     "ignore the multi-dimensional nature of recovery: an athlete with high ACWR but excellent "
     "sleep and HRV may tolerate the load, while an athlete with moderate ACWR but disrupted "
     "sleep may not.")

body(doc,
     "This project reframes the problem: rather than predicting a future injury event, we predict "
     "the athlete\u2019s current training state. We define a composite Grit Score that integrates "
     "training load (via ACWR), autonomic recovery (via HRV z-score relative to personal "
     "baseline), sleep quality (via a composite z-score), resting heart rate trend, body battery, "
     "and stress. A high Grit Score indicates a dangerous overreaching state; a low Grit Score "
     "indicates an undertrained or well-recovered state. We then use this score for three "
     "machine-learning tasks: (1) three-class classification of training state (Overreaching, "
     "Balanced, Undertrained), (2) binary injury-risk classification, and (3) continuous Grit "
     "Score regression. This approach is directly analogous to how WHOOP computes a recovery "
     "score and Garmin computes a training readiness score, but with a transparent, reproducible "
     "methodology.")

# ===== 3. DATASET =====
heading(doc, "3. Dataset", level=1)

body(doc,
     "We use the Synthetic Triathlete Dataset published by Rossi (2025) on Zenodo "
     "(DOI: 10.5281/zenodo.15401061). This dataset was designed for injury prediction research "
     "and contains realistic distributions of training, recovery, and injury patterns across a "
     "population of simulated endurance athletes. It consists of three CSV files:")

body(doc,
     "athletes.csv (1,000 rows): Static athlete profiles including age, sex, VO2max, functional "
     "threshold power (FTP), training experience (years), weekly training hours, and lifestyle "
     "category (sedentary, moderate, active). These represent baseline physiological and "
     "behavioral characteristics that modulate injury risk.")

body(doc,
     "daily_data.csv (366,000 rows): One row per athlete per day over a 366-day period, "
     "containing daily physiological measurements: HRV (ms), resting heart rate (bpm), sleep "
     "hours, deep sleep and REM sleep fractions, sleep quality score, stress score, body battery "
     "morning reading, and a binary injury indicator.")

body(doc,
     "activity_data.csv (384,153 rows): Individual training sessions with duration, Training "
     "Stress Score (TSS), intensity factor, and aerobic/anaerobic training effect. Athletes "
     "may have zero or multiple sessions per day.")

body(doc,
     "After merging the three files and computing rolling features that require 28 days of "
     "warm-up history (for the chronic load denominator in ACWR), we obtain 337,995 usable "
     "athlete-day rows with 24 features. The injury label is imbalanced: 8.14% positive "
     "(29,460 injured days vs. 308,535 healthy days). The load-class distribution, derived from "
     "ACWR tertiles, is 24.6% Undertrained, 50.8% Balanced, and 24.6% Overreaching. To address "
     "class imbalance in the binary injury task, we apply SMOTE (Synthetic Minority "
     "Over-sampling Technique) to the training folds only.")

# ===== 4. METHODS =====
heading(doc, "4. Methods", level=1)

# -- 4.1 Feature Engineering --
heading(doc, "4.1 Feature Engineering", level=2)

body(doc,
     "The core of our approach is the Grit Score, a composite index of training stress designed "
     "so that higher values indicate a more dangerous (overreaching) state and lower values "
     "indicate a safer (undertrained or well-recovered) state. Its components are:")

body(doc,
     "ACWR (Acute-to-Chronic Workload Ratio): The ratio of 7-day rolling sum of TSS to 28-day "
     "rolling sum of TSS, capturing whether recent training load is spiking relative to the "
     "athlete's recent history. Values above 1.3 are considered high-risk in the sports science "
     "literature.")

body(doc,
     "HRV z-score: Each day's HRV expressed as a z-score relative to the athlete's own 14-day "
     "rolling mean and standard deviation. A negative z-score indicates suppressed HRV "
     "(autonomic fatigue); a positive z-score indicates good recovery.")

body(doc,
     "Sleep composite z-score: A combined z-score of sleep hours, deep sleep fraction, REM sleep "
     "fraction, and sleep quality score, again computed relative to personal 14-day baselines. "
     "Lower values indicate poor sleep quality.")

body(doc,
     "RHR 7-day trend: The slope of resting heart rate over the past 7 days. A positive trend "
     "(rising RHR) indicates accumulating fatigue; a negative trend indicates recovery.")

body(doc,
     "Body battery and stress: The morning body battery reading (inversely related to fatigue) "
     "and the daily stress score (directly related to physiological load).")

body(doc,
     "In total, the feature set comprises 24 variables: 5 engineered features (acwr, hrv_zscore, "
     "sleep_composite_z, rhr_trend, grit_score), 13 raw daily signals (hrv, resting_hr, "
     "sleep_hours, deep_sleep, rem_sleep, sleep_quality, stress, body_battery_morning, tss, "
     "duration_minutes, intensity_factor, training_effect_aerobic, training_effect_anaerobic), "
     "and 6 static athlete features (age, vo2max, ftp, training_experience, weekly_training_hours, "
     "gender_enc, lifestyle_enc).")

# Correlation heatmap figure
add_figure(doc,
           os.path.join(PLOTS, "10_correlation_heatmap.png"),
           "Figure 1: Feature correlation heatmap showing relationships among the 24 input features.")

# -- 4.2 Classification Models --
heading(doc, "4.2 Classification Models", level=2)

body(doc,
     "Logistic Regression (scikit-learn): A linear baseline with L2 regularization and "
     "class_weight=balanced to account for class imbalance. This model serves as a reference "
     "point for the discriminability of the feature space.")

body(doc,
     "Random Forest (scikit-learn): An ensemble of 200 decision trees with max_depth=12, "
     "providing nonlinear decision boundaries and built-in feature importance estimates via "
     "mean decrease in impurity.")

body(doc,
     "DNN / MLP (PyTorch): A three-hidden-layer feedforward network with layer sizes "
     "[256, 128, 64], BatchNorm after each hidden layer, ReLU activations, and Dropout(0.3) for "
     "regularization. Trained for 50 epochs with the Adam optimizer (learning rate 1e-3, "
     "batch size 512). For binary injury classification, we use binary cross-entropy loss; for "
     "three-class load classification, we use cross-entropy loss.")

# -- 4.3 Regression Models --
heading(doc, "4.3 Regression Models", level=2)

body(doc,
     "Lasso with Polynomial Features (scikit-learn): Degree-2 polynomial expansion of the "
     "input features to capture pairwise interactions, followed by L1-regularized linear "
     "regression for automatic feature selection and sparsity.")

body(doc,
     "Random Forest (scikit-learn): Same configuration as the classification variant (200 trees, "
     "max_depth=12), with mean squared error as the splitting criterion.")

body(doc,
     "DNN / MLP (PyTorch): Same architecture as the classifier, with a single continuous output "
     "neuron and MSE loss.")

# -- 4.4 Evaluation Strategy --
heading(doc, "4.4 Evaluation Strategy", level=2)

body(doc,
     "All models are evaluated using 5-fold GroupKFold cross-validation, where the grouping "
     "variable is athlete_id. This ensures that all days from a given athlete appear in either "
     "the training set or the test set, but never both, preventing data leakage from temporal "
     "autocorrelation within an athlete's time series.")

body(doc,
     "Classification metrics: ROC-AUC (macro, one-vs-rest), F1-macro, Precision (macro), "
     "Recall (macro), and Accuracy. For the binary injury task, ROC-AUC is the primary metric "
     "because it is threshold-independent and appropriate for imbalanced classes.")

body(doc,
     "Regression metrics: Root Mean Squared Error (RMSE), Mean Absolute Error (MAE), and "
     "coefficient of determination (R\u00b2). RMSE is the primary metric as it penalizes large "
     "errors more heavily, which is important for identifying dangerous training states.")

# ===== 5. RESULTS =====
heading(doc, "5. Results", level=1)

# -- 5.1 Injury Classification --
heading(doc, "5.1 Injury Classification", level=2)

make_table(doc,
           ["Model", "ROC-AUC", "F1-macro", "Precision", "Recall", "Accuracy"],
           [
               ["Logistic Regression", "0.875 \u00b1 0.002", "0.699 \u00b1 0.004",
                "0.663 \u00b1 0.003", "0.810 \u00b1 0.003", "0.858 \u00b1 0.005"],
               ["Random Forest", "0.901 \u00b1 0.003", "0.839 \u00b1 0.003",
                "0.860 \u00b1 0.004", "0.821 \u00b1 0.003", "0.951 \u00b1 0.001"],
               ["DNN (MLP)", "0.948 \u00b1 0.010", "0.871 \u00b1 0.022",
                "0.843 \u00b1 0.024", "0.905 \u00b1 0.017", "0.955 \u00b1 0.008"],
           ],
           col_widths=[1.6, 1.0, 1.0, 1.0, 1.0, 1.0])

cap = doc.add_paragraph("Table 1: Injury classification results (5-fold GroupKFold CV, mean \u00b1 std).")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in cap.runs:
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.italic = True

body(doc,
     "The DNN achieves the highest AUC-ROC (0.948) and F1-macro (0.871), demonstrating that "
     "the nonlinear feature interactions captured by the neural network substantially improve "
     "injury detection. Logistic Regression achieves reasonable recall (0.810) but the lowest "
     "precision (0.663), indicating that it over-predicts injury\u2014likely because the linear "
     "model cannot separate the overlapping regions of the feature space as effectively. Random "
     "Forest offers the best precision-recall balance among non-DNN models, with strong accuracy "
     "(0.951) and AUC (0.901).")

add_figure(doc,
           os.path.join(PLOTS, "01_injury_clf_comparison.png"),
           "Figure 2: Injury classification model comparison\u2014ROC curves and performance metrics.",
           width=Inches(5.2))

# -- 5.2 Load Class Classification --
heading(doc, "5.2 Load Class Classification", level=2)

make_table(doc,
           ["Model", "ROC-AUC", "F1-macro", "Precision", "Recall", "Accuracy"],
           [
               ["Logistic Regression", "0.988 \u00b1 0.000", "0.914 \u00b1 0.001",
                "0.906 \u00b1 0.001", "0.925 \u00b1 0.001", "0.913 \u00b1 0.001"],
               ["Random Forest", "0.987 \u00b1 0.000", "0.914 \u00b1 0.001",
                "0.908 \u00b1 0.001", "0.920 \u00b1 0.001", "0.913 \u00b1 0.001"],
               ["DNN (MLP)", "0.997 \u00b1 0.000", "0.959 \u00b1 0.002",
                "0.960 \u00b1 0.003", "0.958 \u00b1 0.001", "0.959 \u00b1 0.002"],
           ],
           col_widths=[1.6, 1.0, 1.0, 1.0, 1.0, 1.0])

cap = doc.add_paragraph("Table 2: Load class classification results (5-fold GroupKFold CV, mean \u00b1 std).")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in cap.runs:
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.italic = True

body(doc,
     "All three models perform exceptionally well on the three-class load classification task. "
     "Even Logistic Regression achieves an AUC of 0.988, indicating that the ACWR-based class "
     "boundaries are well-separated in the 24-dimensional feature space. The DNN pushes "
     "performance to near-perfect levels, with AUC 0.997 and 95.9% accuracy. The uniformly "
     "small standard deviations across folds confirm that these results generalize reliably "
     "across different athlete subpopulations.")

add_figure(doc,
           os.path.join(PLOTS, "02_load_clf_comparison.png"),
           "Figure 3: Load class classification model comparison\u2014ROC curves and performance metrics.",
           width=Inches(5.2))

# -- 5.3 Grit Score Regression --
heading(doc, "5.3 Grit Score Regression", level=2)

make_table(doc,
           ["Model", "RMSE", "MAE", "R\u00b2"],
           [
               ["Lasso + Poly", "1.809 \u00b1 0.024", "1.372 \u00b1 0.019", "0.973 \u00b1 0.001"],
               ["Random Forest", "1.540 \u00b1 0.030", "1.182 \u00b1 0.020", "0.981 \u00b1 0.001"],
               ["DNN (MLP)", "1.193 \u00b1 0.044", "0.933 \u00b1 0.042", "0.988 \u00b1 0.001"],
           ],
           col_widths=[1.6, 1.5, 1.5, 1.5])

cap = doc.add_paragraph("Table 3: Grit Score regression results (5-fold GroupKFold CV, mean \u00b1 std).")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in cap.runs:
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.italic = True

body(doc,
     "The DNN achieves the best R\u00b2 (0.988) with an RMSE of only 1.19 on a 0\u2013100 scale, "
     "meaning the model's predictions deviate from the true Grit Score by roughly 1.2 points on "
     "average. Even the Lasso baseline with polynomial features achieves R\u00b2 = 0.973, "
     "indicating that the Grit Score is well-predicted by its constituent features even through "
     "a linear lens. All standard deviations are small across folds, confirming stable "
     "generalization.")

add_figure(doc,
           os.path.join(PLOTS, "03_grit_regression_comparison.png"),
           "Figure 4: Grit Score regression model comparison\u2014predicted vs. actual and residual distributions.",
           width=Inches(5.2))

add_figure(doc,
           os.path.join(PLOTS, "07_grit_score_distribution.png"),
           "Figure 5: Distribution of the Grit Score across all athlete-days.",
           width=Inches(4.8))

# -- 5.4 Feature Importance Analysis --
heading(doc, "5.4 Feature Importance Analysis", level=2)

body(doc,
     "We examine feature importance from the Random Forest models, which provide built-in "
     "importance estimates via mean decrease in impurity. The two tasks\u2014injury classification "
     "and Grit Score regression\u2014reveal distinct but complementary importance profiles.")

body(doc,
     "Injury prediction \u2014 top features: (1) deep_sleep at 24.6%, (2) sleep_quality at 12.6%, "
     "(3) sleep_hours at 10.8%, (4) rhr_trend at 8.5%, (5) rem_sleep at 8.3%, and "
     "(6) hrv_zscore at 7.2%. Sleep-related features collectively account for over 55% of total "
     "importance. This finding aligns with sports science literature demonstrating that poor "
     "sleep is among the strongest modifiable risk factors for athletic injury, as it directly "
     "impairs tissue repair, immune function, and neuromuscular coordination.")

add_figure(doc,
           os.path.join(PLOTS, "04_feature_importance_injury.png"),
           "Figure 6: Random Forest feature importance for injury classification.",
           width=Inches(5.0))

body(doc,
     "Grit Score prediction \u2014 top features: (1) hrv_zscore at 49.0%, (2) sleep_composite_z "
     "at 31.9%, (3) body_battery_morning at 15.5%, (4) stress at 2.3%, and (5) acwr at 1.3%. "
     "The Grit Score is primarily driven by HRV deviation from personal baseline and sleep "
     "quality\u2014the two strongest indicators of physiological stress. Notably, ACWR contributes "
     "only 1.3% of importance despite being the traditional load management metric, suggesting "
     "that recovery-side signals carry more information about overall training state than the "
     "load signal alone.")

add_figure(doc,
           os.path.join(PLOTS, "05_feature_importance_grit.png"),
           "Figure 7: Random Forest feature importance for Grit Score regression.",
           width=Inches(5.0))

add_figure(doc,
           os.path.join(PLOTS, "09_injury_vs_features_violin.png"),
           "Figure 8: Violin plots of key features stratified by injury status, illustrating "
           "distributional differences between injured and healthy days.",
           width=Inches(5.5))

# -- 5.5 Hyperparameter Analysis --
heading(doc, "5.5 Hyperparameter Analysis", level=2)

body(doc,
     "A systematic hyperparameter sweep was conducted for all models (results stored in "
     "hp_sweep.pkl). For Logistic Regression, regularization strength C was varied from 0.01 "
     "to 100; performance was stable for C > 0.1, indicating that the feature space is "
     "sufficiently informative that strong regularization is not needed. For Random Forest, "
     "max_depth was varied from 6 to 20; performance plateaued around max_depth=12, with "
     "diminishing returns and slight overfitting beyond that point. For the DNN, we tested "
     "hidden layer sizes of [64, 128, 256] and learning rates of [1e-4, 1e-3, 1e-2]; the "
     "optimal configuration was 256 hidden units per layer with a learning rate of 1e-3. "
     "Lower learning rates converged too slowly within 50 epochs, while higher learning rates "
     "caused training instability.")

# ===== 6. DISCUSSION =====
heading(doc, "6. Discussion", level=1)

body(doc,
     "The DNN consistently achieved the best performance across all three tasks\u2014injury "
     "classification (AUC 0.948), load-class classification (AUC 0.997), and Grit Score "
     "regression (R\u00b2 0.988)\u2014justifying the additional architectural complexity over "
     "simpler baselines. The margin of improvement was largest for injury classification, "
     "where the nonlinear interactions between sleep, HRV, and training load features are most "
     "critical for separating at-risk days from healthy ones.")

body(doc,
     "The dominance of sleep features in injury prediction (accounting for over 55% of Random "
     "Forest importance) is physiologically meaningful and has practical implications. Unlike "
     "training load, which is often considered the primary driver of injury risk, sleep quality "
     "is both a sensitive indicator of accumulated fatigue and a modifiable behavior. Athletes "
     "and coaches could use this insight to prioritize sleep hygiene as an injury prevention "
     "strategy, particularly during periods of high training load.")

body(doc,
     "The Grit Score formulation successfully captures the multi-dimensional nature of training "
     "state, as evidenced by the R\u00b2 of 0.988 and the clean separation of load classes at "
     "AUC 0.997. The score integrates information that no single metric captures alone: an "
     "athlete with high ACWR but excellent HRV and sleep will have a moderate Grit Score, "
     "reflecting that the load is being tolerated. Conversely, moderate ACWR with suppressed "
     "HRV and poor sleep yields a high Grit Score, flagging risk that a load-only metric "
     "would miss.")

body(doc,
     "The GroupKFold cross-validation strategy is critical for valid evaluation. Because daily "
     "observations within an athlete are temporally autocorrelated, a random train/test split "
     "would allow the model to exploit inter-day similarity within the same athlete, inflating "
     "performance estimates. By ensuring that entire athletes are held out, we measure the "
     "model\u2019s ability to generalize to new, unseen individuals\u2014the operationally "
     "relevant scenario.")

body(doc,
     "Several limitations should be acknowledged. First, the dataset is synthetic, and while "
     "it was designed to reflect realistic physiological distributions, it may not capture the "
     "full complexity of real-world data, including measurement noise, missing data, and "
     "behavioral variability. Second, the Grit Score thresholds for the three load classes were "
     "derived from data distribution tertiles rather than clinical validation; in practice, these "
     "thresholds would need to be calibrated against actual injury outcomes. Third, the injury "
     "label in the synthetic data may not perfectly reflect the multi-factorial nature of real "
     "overuse injuries, which can involve biomechanical, nutritional, and psychological factors "
     "not captured in wearable data.")

body(doc,
     "Future work will address these limitations in several directions. We plan to integrate "
     "real wearable data from personal WHOOP, Garmin, and TrainingPeaks accounts using the "
     "respective APIs, enabling validation on actual athlete data. A real-time dashboard with "
     "automatic training plan adjustment based on the Grit Score is under development using a "
     "React frontend. Additionally, we aim to explore temporal modeling architectures (LSTMs, "
     "Transformers) that can explicitly capture the sequential nature of training adaptation, "
     "and to incorporate additional data modalities such as GPS-derived running dynamics and "
     "power meter data from cycling.")

# ===== 7. REFERENCES =====
heading(doc, "7. References", level=1)

refs = [
    "[1] Rossi, L. (2025). Synthetic Triathlete Dataset for Injury Prediction Research (2024). "
    "Zenodo. https://doi.org/10.5281/zenodo.15401061",

    "[2] Rossi, L. (2025). Finding Pre-Injury Patterns in Triathletes from Lifestyle, Recovery "
    "and Load Dynamics Features. arXiv:2511.17610.",

    "[3] Gabbett, T.J. (2016). The training\u2013injury prevention paradox: should athletes be "
    "training smarter and harder? British Journal of Sports Medicine, 50(5), 273\u2013280.",

    "[4] Hulin, B.T., Gabbett, T.J., Blanch, P., Chapman, P., Bailey, D., & Orchard, J.W. "
    "(2016). Spikes in acute workload are associated with increased injury risk in elite "
    "cricket fast bowlers. British Journal of Sports Medicine, 48(8), 708\u2013712.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT)
print(f"Report saved to {OUT}")
